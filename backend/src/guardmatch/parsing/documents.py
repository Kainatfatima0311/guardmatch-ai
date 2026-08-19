"""Getting text out of an uploaded document, or refusing it.

This module exists because of one failure mode, and everything in it is shaped by
that failure mode.

**A scanned PDF has no text layer.** `pypdf` extracts an empty string from it and
reports no error, because nothing went wrong — there simply is no text. If that
empty string were passed on, it would become an *empty CV*, and an empty CV ranks
last. The service would confidently place a candidate at the bottom of a shortlist
because their file could not be read, and the reviewer would see a weak candidate
rather than an unreadable document.

That is precisely the class of silent failure this project is built to prevent, so
extraction refuses rather than returns nothing:

    accepted   .txt, .text, .md      read directly
    accepted   .docx                 python-docx reads the document XML
    accepted   .pdf with a text layer
    refused    .pdf that is scanned or image-only

**OCR was considered and rejected.** Tesseract would let every file through, at
the cost of a very large dependency and, more importantly, mis-read text that
produces the same silent wrong ranking in a new form — a CV whose certifications
were garbled into nothing scores the same as one that could not be read at all,
except now nothing signals it. A refusal a reviewer can act on beats a guess they
cannot check.

Validation is by extension **and** by content. An extension is a claim made by
whoever named the file; the magic bytes are a claim made by the file itself, and a
`.docx` that is really a PDF should not reach a zip parser.
"""

from __future__ import annotations

import io
import zipfile
from typing import Final

import pypdf
from docx import Document

from guardmatch.core.exceptions import ParsingError
from guardmatch.parsing.patterns import MAX_CV_LENGTH

#: A CV is capped at 20,000 characters, so a file far larger than that cannot be
#: one. The allowance is for markup: a .docx is a zip archive and a PDF carries
#: font and layout data, so the file is legitimately much larger than its text.
MAX_UPLOAD_BYTES: Final = 5 * 1024 * 1024

TEXT_EXTENSIONS: Final = frozenset({".txt", ".text", ".md"})
PDF_EXTENSIONS: Final = frozenset({".pdf"})
DOCX_EXTENSIONS: Final = frozenset({".docx"})

SUPPORTED_EXTENSIONS: Final = TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS

#: `.doc` is the old binary Word format, which `python-docx` cannot read at all.
#: Named separately so the refusal can say "convert it" rather than "unsupported".
LEGACY_EXTENSIONS: Final = frozenset({".doc", ".rtf", ".odt", ".pages"})

_PDF_MAGIC: Final = b"%PDF-"
_ZIP_MAGIC: Final = b"PK\x03\x04"


def extension_of(filename: str) -> str:
    """Lower-cased extension including the dot, or "" if there is none."""
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot > 0 else ""


def _reject(message: str) -> ParsingError:
    return ParsingError(message)


def _extract_pdf(payload: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(payload))
    except Exception as exc:
        raise _reject(f"this PDF could not be opened: {exc}") from exc

    if reader.is_encrypted:
        # An encrypted PDF may open with an empty password, but if it does not, the
        # text is genuinely unavailable and pretending otherwise means an empty CV.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise _reject("this PDF is password protected") from exc

    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue

    text = "\n".join(parts)

    if not text.strip():
        # The failure this module exists for. Refused, never returned empty.
        raise _reject(
            "no text layer found — this PDF looks scanned. Paste the text, or "
            "upload a .docx instead."
        )

    return text


def _extract_docx(payload: bytes) -> str:
    try:
        document = Document(io.BytesIO(payload))
    except (zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise _reject(f"this Word file could not be read: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]

    # Tables are where certification lists usually end up in a formatted CV, and
    # `paragraphs` does not include them. Losing them would silently drop the
    # single most important thing this parser looks for.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" ".join(cells))

    text = "\n".join(parts)

    if not text.strip():
        raise _reject("this Word file contains no text")

    return text


def _extract_plain(payload: bytes) -> str:
    # `utf-8-sig` before `utf-8`, and the order is load-bearing. Plain `utf-8`
    # decodes BOM-prefixed bytes successfully, leaving U+FEFF as the first
    # character — which then fails the parser's section-heading pattern, so a CV
    # exported from Windows Notepad would silently lose its first section. Found
    # by a test rather than by a reviewer wondering where PROFILE went.
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            text = payload.decode(encoding)
        except UnicodeDecodeError:
            continue
        if text.strip():
            return text
        raise _reject("this file is empty")

    raise _reject("this file is not readable as text")


def extract_text(filename: str, payload: bytes) -> str:
    """Extract CV text from an uploaded document.

    Args:
        filename: The uploaded name, used for its extension only.
        payload: The file's bytes.

    Returns:
        The extracted text, guaranteed non-empty and within `MAX_CV_LENGTH`.

    Raises:
        ParsingError: The file is too large, of an unsupported kind, not what its
            extension claims, unreadable, empty, or a PDF with no text layer. Every
            case names the file's problem and, where there is one, the way out.
    """
    extension = extension_of(filename)

    if extension in LEGACY_EXTENSIONS:
        raise _reject(
            f"{extension} is an older format this cannot read. Save it as .docx or "
            f".txt and upload again."
        )

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise _reject(f"{extension or 'this file'} is not supported. Accepted: {supported}.")

    if not payload:
        raise _reject("this file is empty")

    if len(payload) > MAX_UPLOAD_BYTES:
        megabytes = len(payload) / (1024 * 1024)
        raise _reject(
            f"{megabytes:.1f} MB is too large for a CV — the limit is "
            f"{MAX_UPLOAD_BYTES // (1024 * 1024)} MB."
        )

    # Content, not just the name. An extension is a claim by whoever named the
    # file; the magic bytes are a claim by the file itself.
    if extension in PDF_EXTENSIONS and not payload.startswith(_PDF_MAGIC):
        raise _reject("this file is named .pdf but its contents are not a PDF")
    if extension in DOCX_EXTENSIONS and not payload.startswith(_ZIP_MAGIC):
        raise _reject("this file is named .docx but its contents are not a Word document")

    if extension in PDF_EXTENSIONS:
        text = _extract_pdf(payload)
    elif extension in DOCX_EXTENSIONS:
        text = _extract_docx(payload)
    else:
        text = _extract_plain(payload)

    if len(text) > MAX_CV_LENGTH:
        raise _reject(
            f"the extracted text is {len(text):,} characters, over the "
            f"{MAX_CV_LENGTH:,} limit by {len(text) - MAX_CV_LENGTH:,}."
        )

    return text
