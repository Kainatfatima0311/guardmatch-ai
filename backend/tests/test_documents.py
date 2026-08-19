"""Document extraction tests.

Extraction is a **new failure surface**, and its failures are quiet. A parser that
returns nothing has not raised anything, and nothing downstream can tell the
difference between "this CV said little" and "this file could not be read". So the
test that matters most here is not that a good file parses — it is that a file
which cannot be read never becomes a candidate.

Fixtures are built in process rather than committed as binaries. A checked-in
scanned PDF would be an opaque blob nobody could verify, and the builders below
state exactly what makes each file the case it is meant to be: the "scanned" PDF
is a real, valid PDF whose only property is having no text layer.
"""

from __future__ import annotations

import io
import zipfile

import pypdf
import pytest
from docx import Document

from guardmatch.core.exceptions import ParsingError
from guardmatch.parsing.documents import (
    MAX_UPLOAD_BYTES,
    extension_of,
    extract_text,
)
from guardmatch.parsing.patterns import MAX_CV_LENGTH

CV_BODY = "PROFILE\nSecurity officer with 6 years of experience."


def blank_pdf() -> bytes:
    """A valid PDF with a page and no text layer — what a scan looks like.

    This is the important fixture. `pypdf` extracts `""` from it and raises
    nothing, because nothing went wrong: there is genuinely no text.
    """
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=595, height=842)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def text_pdf(lines: list[str] | None = None) -> bytes:
    """A minimal PDF carrying a real text layer.

    Hand-assembled rather than produced by a rendering library, so the test does
    not depend on a writer's idea of a page and the bytes stay small enough to read.
    """
    rows = lines or ["PROFILE", "Security officer with 6 years of experience."]
    drawn = " 0 -16 Td ".join(f"({row}) Tj" for row in rows)
    body = f"BT /F1 12 Tf 72 760 Td {drawn} ET".encode()

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(body)).encode() + b" >>\nstream\n" + body + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += str(index).encode() + b" 0 obj\n" + obj + b"\nendobj\n"

    xref_at = len(out)
    out += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode()
        + b" /Root 1 0 R >>\nstartxref\n"
        + str(xref_at).encode()
        + b"\n%%EOF\n"
    )
    return bytes(out)


def docx_file(paragraphs: list[str] | None = None, table: list[str] | None = None) -> bytes:
    document = Document()
    for text in paragraphs or ["PROFILE", "Security officer with 6 years of experience."]:
        document.add_paragraph(text)
    if table:
        row = document.add_table(rows=1, cols=len(table)).rows[0]
        for cell, value in zip(row.cells, table, strict=True):
            cell.text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# The failure this module exists for
# ---------------------------------------------------------------------------


def test_a_pdf_with_no_text_layer_is_refused() -> None:
    """The whole point. A scanned CV must never become an empty candidate.

    `pypdf` returns `""` here and raises nothing. Passing that on would produce an
    empty CV, and an empty CV ranks last — so the service would confidently place a
    candidate at the bottom of a shortlist because their file could not be read,
    and the reviewer would see a weak candidate rather than an unreadable document.
    """
    with pytest.raises(ParsingError, match="no text layer"):
        extract_text("scan.pdf", blank_pdf())


def test_the_refusal_names_a_way_out() -> None:
    """An error a reviewer cannot act on is only marginally better than silence."""
    with pytest.raises(ParsingError) as raised:
        extract_text("scan.pdf", blank_pdf())

    message = str(raised.value)
    assert "scanned" in message
    assert ".docx" in message or "aste" in message


# ---------------------------------------------------------------------------
# What is accepted
# ---------------------------------------------------------------------------


def test_a_pdf_with_a_text_layer_is_read() -> None:
    text = extract_text("cv.pdf", text_pdf())

    assert "PROFILE" in text
    assert "6 years" in text


def test_a_docx_is_read() -> None:
    text = extract_text("cv.docx", docx_file())

    assert "PROFILE" in text
    assert "6 years" in text


def test_docx_tables_are_read_too() -> None:
    """Certification lists live in tables in a formatted CV.

    `python-docx` does not include table cells in `paragraphs`, so reading only
    paragraphs would silently drop the single most important thing this parser
    looks for — and drop it without any signal that something was missed.
    """
    text = extract_text("cv.docx", docx_file(table=["SIA licence", "fire marshal"]))

    assert "SIA licence" in text
    assert "fire marshal" in text


@pytest.mark.parametrize("name", ["cv.txt", "cv.text", "cv.md", "CV.TXT"])
def test_plain_text_is_read_whatever_the_case(name: str) -> None:
    assert "PROFILE" in extract_text(name, CV_BODY.encode())


def test_text_in_an_unusual_encoding_is_read() -> None:
    """A CV exported from an older Windows tool is cp1252, not UTF-8.

    Refusing it would send a reviewer to convert a file that is perfectly readable.
    """
    text = extract_text("cv.txt", "PROFILE\nGuard — 6 years of experience.".encode("cp1252"))

    assert "PROFILE" in text


def test_a_utf8_bom_does_not_leak_into_the_text() -> None:
    text = extract_text("cv.txt", b"\xef\xbb\xbf" + CV_BODY.encode())

    assert text.startswith("PROFILE")


# ---------------------------------------------------------------------------
# What is refused
# ---------------------------------------------------------------------------


def test_an_empty_file_is_refused() -> None:
    with pytest.raises(ParsingError, match="empty"):
        extract_text("cv.txt", b"")


def test_a_whitespace_only_file_is_refused() -> None:
    """Whitespace is not text. Accepting it produces the empty-CV failure again."""
    with pytest.raises(ParsingError, match="empty"):
        extract_text("cv.txt", b"   \n\t  \n")


def test_a_docx_with_no_text_is_refused() -> None:
    with pytest.raises(ParsingError, match="no text"):
        extract_text("cv.docx", docx_file(paragraphs=[""]))


def test_an_oversized_file_is_refused() -> None:
    with pytest.raises(ParsingError, match="too large"):
        extract_text("cv.txt", b"a" * (MAX_UPLOAD_BYTES + 1))


def test_text_over_the_character_limit_is_refused() -> None:
    """The service caps a CV at 20,000 characters, so extraction must too.

    A 60,000-character extraction that passed here would be rejected one step
    later, which reports the problem further from its cause.
    """
    with pytest.raises(ParsingError, match="over the"):
        extract_text("cv.txt", b"a" * (MAX_CV_LENGTH + 1))


def test_a_corrupt_pdf_is_refused() -> None:
    with pytest.raises(ParsingError):
        extract_text("cv.pdf", b"%PDF-1.4\nnot actually a pdf")


def test_a_corrupt_docx_is_refused() -> None:
    with pytest.raises(ParsingError, match="could not be read"):
        extract_text("cv.docx", b"PK\x03\x04garbage")


@pytest.mark.parametrize("name", ["cv.doc", "cv.rtf", "cv.odt", "cv.pages"])
def test_older_formats_are_refused_with_advice(name: str) -> None:
    """Named separately from "unsupported", because the fix is to convert it."""
    with pytest.raises(ParsingError, match="older format"):
        extract_text(name, b"anything")


@pytest.mark.parametrize("name", ["cv.png", "cv.zip", "cv", "cv.exe"])
def test_unsupported_kinds_are_refused_and_list_what_is_accepted(name: str) -> None:
    with pytest.raises(ParsingError, match="not supported"):
        extract_text(name, b"anything")


# ---------------------------------------------------------------------------
# The name is a claim; the bytes are evidence
# ---------------------------------------------------------------------------


def test_a_pdf_that_is_not_a_pdf_is_refused_before_parsing() -> None:
    """An extension is a claim made by whoever named the file.

    Checked against the magic bytes so a mislabelled file is refused with an
    accurate message rather than producing whatever the wrong parser makes of it.
    """
    with pytest.raises(ParsingError, match="not a PDF"):
        extract_text("cv.pdf", docx_file())


def test_a_docx_that_is_not_a_docx_is_refused_before_parsing() -> None:
    with pytest.raises(ParsingError, match="not a Word document"):
        extract_text("cv.docx", text_pdf())


def test_a_zip_bomb_shaped_docx_is_refused_by_size_first() -> None:
    """Size is checked before anything is decompressed."""
    payload = b"PK\x03\x04" + b"\x00" * MAX_UPLOAD_BYTES

    with pytest.raises(ParsingError, match="too large"):
        extract_text("cv.docx", payload)


def test_extraction_never_returns_an_empty_string() -> None:
    """The invariant, asserted across every accepted kind.

    Everything downstream depends on this: a caller that receives text can rank it,
    and a caller that receives an error knows not to. There is no third state where
    something empty flows on looking like a CV.
    """
    for name, payload in [
        ("cv.txt", CV_BODY.encode()),
        ("cv.pdf", text_pdf()),
        ("cv.docx", docx_file()),
    ]:
        assert extract_text(name, payload).strip()


def test_extension_of_handles_awkward_names() -> None:
    assert extension_of("report.2024.final.pdf") == ".pdf"
    assert extension_of("CV.DOCX") == ".docx"
    assert extension_of("noextension") == ""
    assert extension_of(".gitignore") == ""


def test_a_docx_is_a_zip_which_is_why_content_checking_matters() -> None:
    """Documents the reason the magic-byte check for .docx is a zip signature."""
    assert zipfile.is_zipfile(io.BytesIO(docx_file()))
